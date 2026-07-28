"""Export styled investment memos to formatted Word (.docx) documents."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# Repo root / mas_sector_system parent
_PKG_DIR = Path(__file__).resolve().parent
_REPO_ROOT = _PKG_DIR.parent
DEFAULT_OUTPUT_DIR = _REPO_ROOT / "outputs"

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
_BULLET_RE = re.compile(r"^([-*+]|\d+\.)\s+(.+)$")


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def _split_table_row(line: str) -> List[str]:
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def _add_runs(paragraph, text: str) -> None:
    """Add runs with simple **bold** and *italic* markdown support."""
    # Split on **bold** then *italic* within residual pieces
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _style_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def markdown_to_docx(markdown: str, doc: Optional[Document] = None) -> Document:
    """Convert a markdown-ish memo string into a python-docx Document."""
    if doc is None:
        doc = Document()
    _style_document(doc)

    # Reasonable margins for a research memo
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    lines = (markdown or "").replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Markdown heading
        hm = _HEADING_RE.match(stripped)
        if hm:
            level = min(len(hm.group(1)), 3)
            text = hm.group(2).strip()
            # Strip trailing # hashes some models emit
            text = re.sub(r"\s+#+\s*$", "", text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Pipe table: header + separator + body rows
        if _is_table_row(stripped) and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1].strip()):
            header = _split_table_row(stripped)
            i += 2  # skip separator
            rows: List[List[str]] = []
            while i < n and _is_table_row(lines[i].strip()):
                rows.append(_split_table_row(lines[i]))
                i += 1
            cols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Table Grid"
            for c, cell_text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
            for r, row in enumerate(rows):
                for c in range(cols):
                    val = row[c] if c < len(row) else ""
                    cell = table.rows[r + 1].cells[c]
                    cell.text = ""
                    _add_runs(cell.paragraphs[0], val)
            # Spacer after table
            doc.add_paragraph("")
            continue

        # Bullet / numbered list
        bm = _BULLET_RE.match(stripped)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            # Clear default empty run so we control bold/italic ourselves
            if p.runs:
                p.runs[0].text = ""
            _add_runs(p, bm.group(2))
            i += 1
            continue

        # Horizontal rule → skip
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            i += 1
            continue

        # Cover-block style lines (Ticker: / Rating: etc.) as body paragraphs
        # Normal paragraph (may span until blank, but keep line-by-line for short lines)
        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1

    return doc


def export_styled_memo(
    styled_memo: str,
    *,
    ticker: Optional[str] = None,
    output_dir: Optional[Path] = None,
    as_of: Optional[date] = None,
) -> Path:
    """Write styled_memo to outputs/{ticker}_{YYYY-MM-DD}_memo.docx and return the path."""
    out_dir = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    day = as_of or date.today()
    safe_ticker = re.sub(r"[^\w.-]+", "_", (ticker or "SECTOR").strip().upper()) or "SECTOR"
    filename = f"{safe_ticker}_{day.isoformat()}_memo.docx"
    path = out_dir / filename

    doc = markdown_to_docx(styled_memo or "(empty memo)")
    # Title line at top if not already present
    if doc.paragraphs:
        # Insert a subtle subtitle
        pass
    doc.save(str(path))
    return path.resolve()


def docx_export_node(state: dict) -> dict:
    """LangGraph node: export the run's deliverables and print their paths.

    Side effects: write .docx; write clean_memo.json + compliance_audit_log.md;
    finalize cost accounting (console + JSONL); archive to desk memory.

    Output routing is split by audience. The .docx carries thesis content only.
    QC findings, stale XBRL tag warnings, and validation disclosures go to the
    compliance audit log instead of being appended to the memo body — the memo
    gets a one-line pointer so the reader knows the audit exists. Run cost is
    still appended to the memo per the cost-accounting spec, and is repeated in
    the audit log.
    """
    from .artifacts import write_run_artifacts
    from .cost import append_cost_to_memo, finalize_run_cost

    memo = state.get("styled_memo") or ""
    ticker = state.get("ticker")
    if not memo.strip():
        # Fall back so screener-style runs still export something useful
        memo = state.get("final_memo") or ""
    if not memo.strip():
        print("docx export: no styled_memo/final_memo content; skipped.")
        # Still finalize cost so the run is not silent on empty memo.
        cost_update = finalize_run_cost(dict(state))
        # No thesis to ship, but the run still owes a data-quality account.
        try:
            artifact_update = write_run_artifacts(
                {**dict(state), **cost_update},
                context="export_empty_memo",
                write_clean_memo=False,
            )
        except Exception as exc:
            print(f"[artifacts] write failed (non-fatal): {exc}", flush=True)
            artifact_update = {}
        return {**cost_update, **artifact_update}

    qc_status = (state.get("qc_status") or "").upper()

    cost_update = finalize_run_cost(dict(state))
    cost_report = (cost_update.get("cost_report") or state.get("cost_report") or "").strip()

    # Split artifacts first so the memo can point at the audit log by name.
    try:
        artifact_update = write_run_artifacts(
            {**dict(state), **cost_update}, context="export"
        )
    except Exception as exc:
        print(f"[artifacts] write failed (non-fatal): {exc}", flush=True)
        artifact_update = {}

    audit_path = artifact_update.get("compliance_audit_log_path")
    if audit_path:
        memo = (
            memo.rstrip()
            + "\n\nData quality, QC findings, and stale-tag disclosures for this run: "
            + f"`{Path(audit_path).name}`\n"
        )

    memo = append_cost_to_memo(memo, cost_report)

    path = export_styled_memo(memo, ticker=ticker)
    print(f"Saved memo: {path}")
    print(
        f"docx export: qc_status={qc_status or 'n/a'}",
        flush=True,
    )
    # Archive this run for long-term thesis memory (next deep dive on ticker).
    try:
        from .memory import save_run

        save_run({**dict(state), **cost_update, **artifact_update})
    except Exception as exc:
        print(f"[memory] save_run failed (non-fatal): {exc}", flush=True)
    # Return cost + artifact fields so state retains the structured reports.
    return {**cost_update, **artifact_update}
